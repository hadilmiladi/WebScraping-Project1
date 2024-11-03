import requests
from bs4 import BeautifulSoup
from docx import Document

# Get the date from the user
date = input("Please enter the date in the following format MM/DD/YYYY: ")

# Request the page for the specified date
page = requests.get(f"https://www.yallakora.com/match-center/%D9%85%D8%B1%D9%83%D8%B2-%D8%A7%D9%84%D9%85%D8%A8%D8%A7%D8%B1%D9%8A%D8%A7%D8%AA?date={date}")

def main(page):
    # Parse the page content with BeautifulSoup
    source = page.content
    soup = BeautifulSoup(source, "lxml")
    
    # Prepare to store match details
    
    matches_details = []
    championships = soup.find_all("div", {'class': 'matchCard'})

    # Function to extract match information
    def get_match_info(championships):
        championship_title = championships.contents[1].find("h2").text.strip()
        all_matches = championships.contents[3].find_all("div", {'class': 'liItem'})
        
        # Loop through matches
        for match in all_matches:
            team_A = match.find('div', {'class': 'teamA'}).text.strip()
            team_B = match.find('div', {'class': 'teamB'}).text.strip()

            match_score = match.find('div', {'class': 'MResult'}).find_all('span', {'class': 'score'})
            score = f"{match_score[0].text.strip()} - {match_score[1].text.strip()}"
            match_time = match.find('div', {'class': 'MResult'}).find('span', {'class': 'time'}).text.strip()

            # Append match details as a dictionary
            matches_details.append({
                "نوع البطولة": championship_title,
                "الفريق الأول": team_A,
                "الفريق الثاني": team_B,
                "النتيجة": score,
                "التوقيت": match_time
            })

    # Process each championship section
    for championship in championships:
        get_match_info(championship)

    # Print matches_details to check structure
    print(matches_details)

    # Create a Word document
    doc = Document()
    doc.add_heading('Match Details', 0)

    # Add a table to the Word document
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    # Define column headers in Arabic
    headers = ["نوع البطولة", "الفريق الأول", "الفريق الثاني", "النتيجة", "التوقيت"]
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    # Add rows to the table, with key checking
    for match in matches_details:
        row_cells = table.add_row().cells
        for idx, key in enumerate(headers):
            row_cells[idx].text = match.get(key, "N/A")  # Default to "N/A" if key is missing

    # Save the document
    output_file_path = '/users/hadil/documents/scrap/matches-details.docx'
    doc.save(output_file_path)
    print("Word document created successfully at", output_file_path)

# Run the main function
main(page)
